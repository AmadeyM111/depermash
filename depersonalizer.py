"""
Depersonalizer (PII).

ALgorithms:  NER (Natasha + Presidio/spaCy/stanza), Regex,
            K-anonymity, Differential privacy (Laplas).
Formats:    PDF (pymupdf/fitz), DOCX (python-docx), XLSX (openpyxl),
            CSV (pandas), PNG/JPG (Tesseract OCR), TXT/plain.
            Structured: JSON

Deploy:
    pip install -r requirements.txt
    python -m spacy download ru_core_news_sm
    python -m spacy download en_core_web_sm
    # stanza (optional):
    python -c "import stanza; stanza.download('ru')"

Pipeline:
    Extract -> Detect (assambly Regex+NER) -> Anonymize -> Rebuild
"""

from __future__ import annotations

import hashlib
import warnings
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterator

import numpy as np

# ──────────────────────────────────────────────────────────────
# Optional dependecies — graceful degradation
# ──────────────────────────────────────────────────────────────
try:
    from natasha import Doc as NatashaDoc
    from natasha import MorphVocab, NewsEmbedding, NewsNERTagger, Segmenter  # noqa: F401
    _NATASHA_OK = True
except ImportError:
    warnings.warn("natasha не установлена — RU-NER отключён.", stacklevel=1)
    _NATASHA_OK = False

try:
    from presidio_analyzer import AnalyzerEngine, RecognizerResult
    from presidio_analyzer.nlp_engine import NlpEngineProvider
    from presidio_anonymizer import AnonymizerEngine          # noqa: F401
    from presidio_anonymizer.entities import OperatorConfig   # noqa: F401
    _PRESIDIO_OK = True
except ImportError:
    warnings.warn("presidio не установлена — мультиязычный NER отключён.", stacklevel=1)
    _PRESIDIO_OK = False

try:
    import stanza  # noqa: F401
    from presidio_analyzer.nlp_engine import StanzaNlpEngine
    _STANZA_OK = True
except ImportError:
    _STANZA_OK = False

try:
    import fitz  # pymupdf
    _PYMUPDF_OK = True
except ImportError:
    warnings.warn("pymupdf не установлен — обработка PDF отключена.", stacklevel=1)
    _PYMUPDF_OK = False

try:
    from docx import Document
    from docx.oxml.ns import qn  # noqa: F401
    _DOCX_OK = True
except ImportError:
    warnings.warn("python-docx не установлен — обработка DOCX отключена.", stacklevel=1)
    _DOCX_OK = False

try:
    import pandas as pd
    _PANDAS_OK = True
except ImportError:
    _PANDAS_OK = False

try:
    from openpyxl import Workbook, load_workbook
    _OPENPYXL_OK = True
except ImportError:
    warnings.warn("openpyxl не установлен — обработка XLSX отключена.", stacklevel=1)
    _OPENPYXL_OK = False

try:
    from PIL import Image, ImageDraw, ImageFont
    import pytesseract
    _PILLOW_OK = True
    _TESSERACT_OK = True
except ImportError:
    warnings.warn("Pillow или pytesseract не установлены — обработка изображений отключена.", stacklevel=1)
    _PILLOW_OK = False
    _TESSERACT_OK = False

import re


# ══════════════════════════════════════════════════════════════
# 1.  ДАННЫЕ И КОНСТАНТЫ
# ══════════════════════════════════════════════════════════════

# Маппинг типов сущностей -> рус. метки для плейсхолдеров
ENTITY_LABEL_RU: dict[str, str] = {
    "PER":           "ФИО",
    "PERSON":        "ФИО",
    "LOC":           "АДРЕС",
    "LOCATION":      "АДРЕС",
    "ORG":           "ОРГ",
    "ORGANIZATION":  "ОРГ",
    "PHONE":         "ТЕЛЕФОН",
    "PHONE_NUMBER":  "ТЕЛЕФОН",
    "EMAIL":         "EMAIL",
    "EMAIL_ADDRESS": "EMAIL",
    "PASSPORT":      "ПАСПОРТ",
    "SNILS":         "СНИЛС",
    "INN":           "ИНН",
    "DATE_TIME":     "ДАТА",
    "DATE":          "ДАТА",
    "AGE":           "ВОЗРАСТ",
    "CREDIT_CARD":   "КАРТА",
    "IBAN_CODE":        "IBAN",
    "IP_ADDRESS":       "IP",
    "URL":              "URL",
    "DOCUMENT_NUMBER":  "ДОКУМЕНТ_№",
    "USER_ID":          "ID_ПОЛЬЗОВАТЕЛЯ",
    "CLIENT_ID":        "ID_КЛИЕНТА",
    # Специальные категории (ст. 10 ФЗ-152)
    "HEALTH":           "ЗДОРОВЬЕ",
    "CRIMINAL":         "СУДИМОСТЬ",
    "RELIGION":         "РЕЛИГИЯ",
    "POLITICS":         "ПОЛИТИКА",
    "NATIONALITY":      "НАЦИОНАЛЬНОСТЬ",
    # Банковская тайна
    "INCOME":           "ДОХОД",
    "CREDIT_HISTORY":   "КРЕДИТНАЯ_ИСТОРИЯ",
    "TRANSACTION_ID":   "ID_СДЕЛКИ",
    "ACCOUNT":          "СЧЕТ",
    # Биометрия
    "BIOMETRIC":        "БИОМЕТРИЯ",
    "FILE_PHOTO":       "ФОТО",
    "FILE_VOICE":       "ГОЛОС",
    # Дополнительные документы
    "DRIVER_LICENSE":   "ВОДИТЕЛЬСКОЕ_УДОСТОВЕРЕНИЕ",
    "MILITARY_ID":      "ВОЕННЫЙ_БИЛЕТ",
    "INSURANCE_POLICY": "ПОЛИС",
}

# Месяцы для regex дат прописью
_MONTHS_RU = (
    "января|февраля|марта|апреля|мая|июня|"
    "июля|августа|сентября|октября|ноября|декабря"
)

# Regex-паттерны для российских PII (специфичность убывает сверху вниз)
_RU_REGEX_PATTERNS: list[tuple[str, str]] = [
    # Адрес после «по адресу:» — до «;» или конца строки
    ("LOCATION",
     r"(?:зарегистрированн(?:ый|ая|ое|ую)\s+по\s+адресу\s*:\s*)"
     r"([^;\n]+)"),
    # Адрес в формате «ул./пр./пер. Название, дом X, ...кв. Y»
    ("LOCATION",
     r"(?:ул\.|пр\.|пер\.|просп\.|бул\.|наб\.|ш\.)\s*[А-ЯЁа-яё\w\s\-]+,"
     r"\s*(?:дом|д\.)\s*\d[\d\w\s,./лит корпАа-яё]*?(?:кв\.\s*\d+)"),
    # СНИЛС: 123-456-789 01 | 123 456 789 01 | 018-070-39 (сокращённый)
    ("SNILS",
     r"\b\d{3}[-\s]\d{3}[-\s]\d{2,3}(?:[-\s]\d{2})?\b"),
    # ИНН: 10 (юр. лицо) или 12 цифр (физ. лицо)
    ("INN",
     r"(?i)(?:инн\s*:?\s*)?\b\d{10}(?:\d{2})?\b"),
    # Серия + номер паспорта РФ: 4510 123456 | 45 10 123456
    ("PASSPORT",
     r"\b\d{4}\s\d{6}\b|\b\d{2}\s\d{2}\s\d{6}\b"),
    # Код подразделения: 780-189
    ("PASSPORT",
     r"(?:код\s+подразделения\s*)(\d{3}[-‐]\d{3})"),
    # «выдан ... отделом/отделением/управлением ...» — контекст выдачи документа
    ("LOCATION",
     r"(?:выдан[а]?\s+)(\d{1,3}\s+)?(?:отдел(?:ом|ением)?|управлением|ГУ|ОУФМС|УФМС|ТП)\s+"
     r"[^,;.\n]{5,80}"),
    # Номер доверенности / реестровый номер: 78/670-н/78-2025-4-392
    ("DOCUMENT_NUMBER",
     r"\b(\d{1,4}/\d[\d\w\-/нН]+)"),
    # Банковская карта (4 группы по 4 цифры)
    ("CREDIT_CARD",
     r"\b(?:\d{4}[\s\-]){3}\d{4}\b"),
    # Российский телефон: +7/8/7 + различные форматы
    ("PHONE_NUMBER",
     r"(?:\+7|8|7)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}"),
    # Email
    ("EMAIL_ADDRESS",
     r"[a-zA-Z0-9_.+\-]+@[a-zA-Z0-9\-]+\.[a-zA-Z0-9.\-]+"),
    # URL (до email-пат., чтобы домены не мешали)
    ("URL",
     r"https?://[^\s\"'<>]+"),
    # Дата прописью: «12 мая 1959 года» | «12 мая 1959 г.»
    ("DATE_TIME",
     r"\b\d{1,2}\s+(?:" + _MONTHS_RU + r")\s+\d{4}\s*(?:г(?:ода|\.)?)?"),
    # Дата: DD.MM.YYYY | DD/MM/YYYY | YYYY-MM-DD
    ("DATE_TIME",
     r"\b(?:\d{2}[./]\d{2}[./]\d{4}|\d{4}-\d{2}-\d{2})\b"),
    # Дата/время ISO 8601: 2026-04-12T15:20:00
    ("DATE_TIME",
     r"\b\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}(?:\.\d+)?(?:Z|[+-]\d{2}:\d{2})?\b"),
    # ID пользователя/клиента: "user_id": 6192 (JSON формат)
    ("USER_ID",
     r'(?:user_id|client_id)\s*:\s*(\d{4,})'),
    # IP-адрес v4
    ("IP_ADDRESS",
     r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),
    # IBAN
    ("IBAN_CODE",
     r"\b[A-Z]{2}\d{2}[A-Z0-9]{4,30}\b"),
    # Русские отчества: Николаевна, Александрович, и т.п.
    ("PERSON",
     r"\b[А-ЯЁ][а-яё]+(?:ович|евич|ич|овна|евна|ична|инична)\b"),
    # Русские/украинские фамилии (типичные суффиксы) перед именем/отчеством
    ("PERSON",
     r"\b[А-ЯЁ][а-яё]+(?:ов|ев|ёв|ин|ын|ский|ская|цкий|цкая|ова|ева|ёва|ина|ына|енко|ейко|юк|чук)"
     r"(?:\s+[А-ЯЁ][а-яё]+){1,2}\b"),

    # ──────────────────────────────────────────────────────────────
    # Специальные категории (ст. 10 ФЗ-152)
    # ──────────────────────────────────────────────────────────────

    # Здоровье: диагнозы, болезни, инвалидность
    ("HEALTH",
     r"(?i)(?:диагноз?:?\s*[:\"]?\s*|(?:сахарный\s+)?диабет|инвалидность|онколог|гипертония|астма|вирус\s*\w*|заболевание)"
     r"|(?:группа\s+инвалидности|класс\s+фк|инфаркт|инсульт)"),

    # Судимость: упоминания УК РФ, судимости, статей
    ("CRIMINAL",
     r"(?i)(?:судимость|(?:имееется\s+)?(?:уголовная\s+)?ответственность|осужден"
     r"|\s(?:ст\.|статья\s*)\s*\d+\s+УК\s*РФ|привлечение\s+к\s+ответственности"
     r"|(?:уголовное\s+)?дело|суд(?:имость)?\s*(?:за|от|в)?)"),

    # Религия: упоминания религий, конфессий
    ("RELIGION",
     r"(?i)(?:православн(?:ый|ая|ое)|мусульманин(?:ский|ка)?|католик|протестант|иудей|буддист"
     r"|мусульманский|христианск(?:ий|ая)|атеист|религиозн(?:ый|ая|ое)"
     r"|(?:прихожанин|член)\s+(?:церкви|храма|мечети|синагоги))"),

    # Политика: партии, политические взгляды
    ("POLITICS",
     r"(?i)(?:член\s+партии|партия\s+|(?:политические)?\s*взгляды"
     r"|(?:КПРФ|LDPR|Справедливая\s+Россия|Единая\s+Россия|Яблоко)"
     r"|(?:депутат|активист|протестовавший?))"),

    # Национальность: упоминания национальностей, народностей
    ("NATIONALITY",
     r"(?i)(?:русск(?:ий|ая|кое)|украинск(?:ий|ая|кое)|белорус|казах|татарин|армян(?:ин|е)"
     r"|(?:национальность|этнос|народность)\s*[:\"]?\s*[А-ЯЁа-яё\-]+)"),

    # ──────────────────────────────────────────────────────────────
    # Банковская тайна
    # ──────────────────────────────────────────────────────────────

    # Доходы, зарплаты, налоги
    ("INCOME",
     r"(?i)(?:доход(?:ы|ах)?|зарплата|заработн(?:ая|ой)\s+плата"
     r"|недвижимость|недвиж(?:имое)?|капитал|сбережения|активы)"
     r"|(?:НДФЛ|НДС|налог)\s*(?:\:|\"|]|\s+)\s*\d{3,}"),

    # Кредитная история: просрочки, долги
    ("CREDIT_HISTORY",
     r"(?i)(?:кредитн(?:ая|ой)?\s+истор(?:ия|ии)?|просрочк(?:а|и)?|задолженност(?:ь|и)"
     r"|долгов(?:ая|ой)?\s+нагрузк(?:а|и)?|кредитный\s+рейтинг"
     r"|(?:плохая|хорошая)\s+кредитн(?:ая|ая)\s+истор(?:ия|ии)?)"),

    # ID сделок, УИд, QR-коды
    ("TRANSACTION_ID",
     r"(?i)(?:уид\s+сделки|ид\s+сделки|qr[-\s]*код|штрих[-\s]*код"
     r"|(?:уникальный\s+)?идентификатор\s+сделки"
     r"|qr[-:\s]*\d{4,}[-:\s]*\d{4,})"),

    # Банковские счета: российский формат (20 цифр, начинается с 40817...)
    ("ACCOUNT",
     r"\b(?:40817\d{13}|\\d{20})\b"),

    # Пароли, ключи (для банковской безопасности)
    ("ACCOUNT",  # Используем ACCOUNT вместо PASSWORD для простоты
     r"(?i)(?:пароль|ключ|пин-код|код\s+подтверждения)\s*[:\"]?\s*\S+"),

    # ──────────────────────────────────────────────────────────────
    # Биометрия
    # ──────────────────────────────────────────────────────────────

    # Фотографии: упоминания файлов изображений
    ("FILE_PHOTO",
     r"(?i)(?:фотограф(?:ия|ии)? из?\s*(?:паспорта|анкеты|документа)"
     r"|(?:scan|скан)\s*(?:страницы?|стр\.?\s*\d+)"
     r"|photo[-_\s]*\d+|фото[-_\s]*страница[-_\s]*\d+)"),

    # Записи голоса, аудиофайлы
    ("FILE_VOICE",
     r"(?i)(?:запись\s+голоса|аудиозапись|голосовое\s+сообщение"
     r"|call[-_\s]*recording|录音|voice[-_\s]*note"
     r"|\.(?:mp3|wav|ogg|m4a)(?:\?|\s|$))"),

    # ДНК, биометрические данные
    ("BIOMETRIC",
     r"(?i)(?:днк[-_\s]?профиль|биометрически(?:ий|ая|ое)?\s*данные"
     r"|радужн(?:ая|ой)?\s+глаз(?:а)?|отпечаток\s+пальц(?:а|ев)?"
     r"|iris[-_\s]scan|fingerprint[-_\s]scan)"),

    # ──────────────────────────────────────────────────────────────
    # Дополнительные документы
    # ──────────────────────────────────────────────────────────────

    # Водительское удостоверение
    ("DRIVER_LICENSE",
     r"(?i)(?:водительск(?:ое|ая|ое)?\s+удостоверени(?:е|я|ия)?"
     r"|прав(?:а|ах)?\s+(?:на\s+управление|водительск)"
     r"|(?:д(?:р)?\s*)?[А-ЯЁA-Z]{2}\s*\d{3,6})"),

    # Военный билет
    ("MILITARY_ID",
     r"(?i)(?:военн(?:ый|ая|ое)?\s+билет(?:у|е|ы)?"
     r"|мк[-_\s]?\d{5,}|учетн(?:ый|ая|ое)?\s+(?:военн(?:ый|ая)?)?\s*документ)"),

    # Полисы ОМС/ДМС
    ("INSURANCE_POLICY",
     r"(?i)(?:полис(?:\s+(?:омс|дмс|медицинск|страхов))?"
     r"|страхов(?:ое|ой|ая|ые)?\s+полис"
     r"|медицинск(?:ая|ое|ие)?\s+страховк(?:а|и)"
     r"|един(?:ый|ое|ые)\s+полис)"),

    # Улучшенный паттерн для голосовых файлов (catch filenames)
    ("FILE_VOICE",
     r"(?i)(?:запись\s+голоса|аудиозапись|голосовое\s+сообщение"
     r"|call[-_\s]*recording|录音|voice[-_\s]note"
     r"|\.(?:mp3|wav|ogg|m4a)(?:\?|\s|$)"
     r"|[-_\s]voice[-_\s]|[-_\s]запись[-_\s]|\brecord_\w+)"),
]


@dataclass
class PIIMatch:
    """Одна найденная PII-сущность в тексте."""
    start:         int
    end:           int
    entity_type:   str          # PHONE_NUMBER | PER | PASSPORT | …
    original_text: str
    confidence:    float = 1.0
    source:        str   = ""   # "regex" | "natasha" | "presidio"

    def __len__(self) -> int:
        return self.end - self.start


# ══════════════════════════════════════════════════════════════
# 2.  ТРЕКЕР СУЩНОСТЕЙ  (последовательные ID)
# ══════════════════════════════════════════════════════════════

class EntityTracker:
    """
    Запоминает оригинальный текст → плейсхолдер.
    Одинаковый оригинал всегда получает один и тот же плейсхолдер.
    """

    def __init__(self) -> None:
        self._map:     dict[str, str] = {}
        self._counter: dict[str, int] = defaultdict(int)

    def get_placeholder(self, original: str, entity_type: str) -> str:
        if original not in self._map:
            label = ENTITY_LABEL_RU.get(entity_type, entity_type.upper())
            self._counter[label] += 1
            self._map[original] = f"[{label}_{self._counter[label]}]"
        return self._map[original]

    @staticmethod
    def get_hash(original: str) -> str:
        digest = hashlib.sha256(original.encode()).hexdigest()[:8].upper()
        return f"[HASH_{digest}]"

    def mapping(self) -> dict[str, str]:
        return dict(self._map)


# ══════════════════════════════════════════════════════════════
# § 3.  ДЕТЕКТОРЫ
# ══════════════════════════════════════════════════════════════

class RegexDetector:
    """Regex-детектор для российских PII-паттернов."""

    def __init__(self) -> None:
        self._compiled = [
            (etype, re.compile(pattern))
            for etype, pattern in _RU_REGEX_PATTERNS
        ]

    def detect(self, text: str) -> list[PIIMatch]:
        matches: list[PIIMatch] = []
        for entity_type, rx in self._compiled:
            for m in rx.finditer(text):
                # Если есть группа захвата — используем её (контекст не заменяется)
                if m.lastindex:
                    start, end = m.start(m.lastindex), m.end(m.lastindex)
                    original = m.group(m.lastindex)
                else:
                    start, end = m.start(), m.end()
                    original = m.group()
                matches.append(PIIMatch(
                    start=start,
                    end=end,
                    entity_type=entity_type,
                    original_text=original,
                    confidence=0.95,
                    source="regex",
                ))
        return matches


class NatashaDetector:
    """NER для русского языка через Natasha (PER / LOC / ORG)."""

    def __init__(self) -> None:
        if not _NATASHA_OK:
            raise RuntimeError("natasha не установлена: pip install natasha")
        self._segmenter  = Segmenter()
        self._emb        = NewsEmbedding()
        self._ner_tagger = NewsNERTagger(self._emb)

    def detect(self, text: str) -> list[PIIMatch]:
        doc = NatashaDoc(text)
        doc.segment(self._segmenter)
        doc.tag_ner(self._ner_tagger)
        return [
            PIIMatch(
                start=span.start,
                end=span.stop,
                entity_type=span.type,
                original_text=span.text,
                confidence=0.85,
                source="natasha",
            )
            for span in doc.spans
        ]


class PresidioDetector:
    """
    Мультиязычный NER через Presidio.

    NLP-бэкенды:
      - spaCy (по умолчанию):  ru_core_news_sm / en_core_web_sm
      - stanza (use_stanza=True): требует stanza + скачанную ru-модель
    """

    def __init__(self, language: str = "ru", use_stanza: bool = False) -> None:
        if not _PRESIDIO_OK:
            raise RuntimeError(
                "presidio не установлена: "
                "pip install presidio-analyzer presidio-anonymizer"
            )
        self._language = language
        self._analyzer = self._build_analyzer(language, use_stanza)

    @staticmethod
    def _build_analyzer(lang: str, use_stanza: bool) -> "AnalyzerEngine":
        if use_stanza and _STANZA_OK:
            nlp_engine = StanzaNlpEngine(
                models=[{"lang_code": lang, "model_name": lang}]
            )
            return AnalyzerEngine(
                nlp_engine=nlp_engine,
                supported_languages=[lang],
            )
        models = [
            {"lang_code": "ru", "model_name": "ru_core_news_sm"},
            {"lang_code": "en", "model_name": "en_core_web_sm"},
        ]
        cfg = {"nlp_engine_name": "spacy", "models": models}
        try:
            provider   = NlpEngineProvider(nlp_configuration=cfg)
            nlp_engine = provider.create_engine()
            return AnalyzerEngine(
                nlp_engine=nlp_engine,
                supported_languages=["ru", "en"],
            )
        except Exception as exc:
            raise RuntimeError(
                f"Presidio: не удалось загрузить spaCy модели ({exc}).\n"
                "Запустите: python -m spacy download ru_core_news_sm"
            ) from exc

    def detect(self, text: str) -> list[PIIMatch]:
        results: list["RecognizerResult"] = self._analyzer.analyze(
            text=text,
            language=self._language,
            entities=list(ENTITY_LABEL_RU.keys()),
        )
        return [
            PIIMatch(
                start=r.start,
                end=r.end,
                entity_type=r.entity_type,
                original_text=text[r.start:r.end],
                confidence=r.score,
                source="presidio",
            )
            for r in results
        ]


# ══════════════════════════════════════════════════════════════
# § 4.  АНСАМБЛЬ — объединение и дедупликация spans
# ══════════════════════════════════════════════════════════════

class EnsembleDetector:
    """
    Запускает все включённые детекторы, объединяет результаты.
    При перекрытии спанов побеждает тот, у которого выше confidence;
    при равных — тот, который длиннее.
    """

    def __init__(
        self,
        use_regex:    bool = True,
        use_natasha:  bool = True,
        use_presidio: bool = True,
        use_stanza:   bool = False,
        language:     str  = "ru",
    ) -> None:
        self._detectors: list[Any] = []
        if use_regex:
            self._detectors.append(RegexDetector())
        if use_natasha and _NATASHA_OK:
            self._detectors.append(NatashaDetector())
        if use_presidio and _PRESIDIO_OK:
            try:
                self._detectors.append(PresidioDetector(language, use_stanza))
            except RuntimeError as exc:
                warnings.warn(str(exc), stacklevel=2)

    def detect(self, text: str) -> list[PIIMatch]:
        raw: list[PIIMatch] = []
        for det in self._detectors:
            raw.extend(det.detect(text))
        return self._merge_spans(raw)

    @staticmethod
    def _merge_spans(matches: list[PIIMatch]) -> list[PIIMatch]:
        if not matches:
            return []
        sorted_m = sorted(matches, key=lambda m: (m.start, -(m.end - m.start)))
        merged:  list[PIIMatch] = []
        current = sorted_m[0]
        for nxt in sorted_m[1:]:
            if nxt.start < current.end:   # перекрытие
                current = max(
                    current, nxt,
                    key=lambda m: (m.confidence, len(m))
                )
            else:
                merged.append(current)
                current = nxt
        merged.append(current)
        return merged


# ══════════════════════════════════════════════════════════════
# § 5.  АНОНИМИЗАТОР ТЕКСТА
# ══════════════════════════════════════════════════════════════

class TextAnonymizer:
    """
    Заменяет PII-spans в тексте на маски/плейсхолдеры/хэши.
    Замена идёт справа налево, чтобы не смещались индексы.
    """

    def __init__(
        self,
        mode:    str                 = "placeholder",
        tracker: EntityTracker | None = None,
    ) -> None:
        assert mode in ("placeholder", "hash", "mask"), \
            "mode должен быть 'placeholder', 'hash' или 'mask'"
        self.mode    = mode
        self.tracker = tracker or EntityTracker()

    def anonymize(self, text: str, matches: list[PIIMatch]) -> str:
        for m in sorted(matches, key=lambda m: m.start, reverse=True):
            text = text[:m.start] + self._replacement(m) + text[m.end:]
        return text

    def _replacement(self, m: PIIMatch) -> str:
        if self.mode == "placeholder":
            return self.tracker.get_placeholder(m.original_text, m.entity_type)
        if self.mode == "hash":
            return self.tracker.get_hash(m.original_text)
        return "█" * len(m.original_text)   # mask


# ══════════════════════════════════════════════════════════════
# § 6.  K-АНОНИМНОСТЬ  (табличные данные / pandas DataFrame)
# ══════════════════════════════════════════════════════════════

class KAnonymityProcessor:
    """
    Обобщает квази-идентификаторы в pandas DataFrame так, чтобы
    каждая уникальная комбинация QI встречалась ≥ k раз.

    Стратегии обобщения (передаются через `strategies`):
      "age_range"  — возраст → диапазон «20-29»
      "year_only"  — дата → только год
      "zip_prefix" — почтовый индекс → первые 3 символа + «***»
      "suppress"   — удалить записи, нарушающие k-условие

    Пример:
        kp = KAnonymityProcessor(k=3)
        df_anon = kp.process(df, quasi_ids=["age", "zip", "birth_date"])
    """

    def __init__(self, k: int = 3) -> None:
        if not _PANDAS_OK:
            raise RuntimeError("pandas не установлен: pip install pandas")
        self.k = k

    def process(
        self,
        df:          "pd.DataFrame",
        quasi_ids:   list[str],
        strategies:  dict[str, str] | None = None,
    ) -> "pd.DataFrame":
        df    = df.copy()
        strat = strategies or {}
        for col in quasi_ids:
            s      = strat.get(col, self._guess_strategy(df, col))
            df[col] = self._generalize(df[col], s)
        # Suppress: удаляем группы меньше k
        counts     = df.groupby(quasi_ids)[quasi_ids[0]].transform("count")
        result     = df[counts >= self.k].reset_index(drop=True)
        n_dropped  = len(df) - len(result)
        if n_dropped:
            warnings.warn(
                f"K-анонимность: подавлено {n_dropped} записей (< k={self.k}).",
                stacklevel=2,
            )
        return result

    @staticmethod
    def _guess_strategy(df: "pd.DataFrame", col: str) -> str:
        sample = df[col].dropna().head(50)
        if pd.api.types.is_numeric_dtype(sample):
            return "age_range"
        try:
            pd.to_datetime(sample, dayfirst=True)
            return "year_only"
        except Exception:
            return "zip_prefix"

    @staticmethod
    def _generalize(series: "pd.Series", strategy: str) -> "pd.Series":
        if strategy == "age_range":
            def _r(v: Any) -> str:
                try:
                    a = int(float(v)); lo = (a // 10) * 10
                    return f"{lo}-{lo + 9}"
                except (ValueError, TypeError):
                    return str(v)
            return series.apply(_r)
        if strategy == "year_only":
            def _y(v: Any) -> str:
                try:
                    return str(pd.to_datetime(v, dayfirst=True).year)
                except Exception:
                    return str(v)
            return series.apply(_y)
        if strategy == "zip_prefix":
            return series.astype(str).str[:3] + "***"
        return series


# ══════════════════════════════════════════════════════════════
# § 7.  ДИФФЕРЕНЦИАЛЬНАЯ ПРИВАТНОСТЬ  (механизм Лапласа)
# ══════════════════════════════════════════════════════════════

class DifferentialPrivacyProcessor:
    """
    Добавляет Лапласовский шум к числовым столбцам DataFrame.

    Параметры:
        epsilon (ε): бюджет приватности (↓ epsilon = ↑ шум/приватность).
                     Рекомендуемый диапазон: 0.1 – 1.0.
        sensitivity: чувствительность функции запроса.
                     Для нормализованных данных: 1.0.
                     Для зарплат: укажите диапазон (max−min).

    Шум: Laplace(0, sensitivity / epsilon)
    Стандарт: https://www.microsoft.com/en-us/research/publication/
              the-algorithmic-foundations-of-differential-privacy/
    """

    def __init__(self, epsilon: float = 1.0, sensitivity: float = 1.0) -> None:
        assert epsilon > 0, "epsilon должен быть > 0"
        self.epsilon     = epsilon
        self.sensitivity = sensitivity
        self._rng        = np.random.default_rng()

    def add_noise(
        self,
        df:              "pd.DataFrame",
        numeric_columns: list[str],
        round_to:        int | None = 0,
    ) -> "pd.DataFrame":
        """Добавляет Лапласовский шум к числовым столбцам."""
        if not _PANDAS_OK:
            raise RuntimeError("pandas не установлен.")
        df    = df.copy()
        scale = self.sensitivity / self.epsilon
        for col in numeric_columns:
            if col not in df.columns:
                continue
            noise = self._rng.laplace(loc=0.0, scale=scale, size=len(df))
            df[col] = df[col].astype(float) + noise
            if round_to is not None:
                df[col] = df[col].round(round_to)
        return df

    def privatize_count(self, count: int) -> int:
        """Приватизирует скалярный счётчик."""
        noise = self._rng.laplace(0.0, self.sensitivity / self.epsilon)
        return max(0, round(count + noise))


# ══════════════════════════════════════════════════════════════
# § 8.  ПРОЦЕССОРЫ ФАЙЛОВ
# ══════════════════════════════════════════════════════════════

class PDFProcessor:
    """
    Extract/Rebuild для PDF через pymupdf (fitz).

    Rebuild: физическое зачернение + вставка текста-плейсхолдера
             через page.add_redact_annot(fill_text=...).
    """

    def __init__(self) -> None:
        if not _PYMUPDF_OK:
            raise RuntimeError("pymupdf не установлен: pip install pymupdf")

    def extract_text(self, path: str | Path) -> str:
        parts: list[str] = []
        with fitz.open(str(path)) as doc:
            for page in doc:
                parts.append(page.get_text())
        return "\n".join(parts)

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],
        fill_color:   tuple[float, ...] = (1.0, 1.0, 1.0),
        font_size:    float = 8,
    ) -> None:
        with fitz.open(str(input_path)) as doc:
            for page in doc:
                for original, placeholder in replacements.items():
                    for quad in page.search_for(original, quads=True):
                        page.add_redact_annot(
                            quad=quad,
                            fill=fill_color,
                            text=placeholder,
                            fontsize=font_size,
                        )
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
            doc.save(str(output_path), garbage=4, deflate=True)


class DocxProcessor:
    """
    Extract/Rebuild для DOCX через python-docx.

    Rebuild: замена в runs параграфов + поколонная замена в таблицах
             (каждый столбец обрабатывается отдельно — per task spec).
    """

    def __init__(self) -> None:
        if not _DOCX_OK:
            raise RuntimeError("python-docx не установлен: pip install python-docx")

    def extract_text(self, path: str | Path) -> str:
        doc   = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],
    ) -> None:
        doc = Document(str(input_path))

        # 1. Обычные параграфы + header/footer
        for para in self._iter_all_paragraphs(doc):
            self._replace_in_paragraph(para, replacements)

        # 2. Таблицы: каждый столбец — отдельно
        for table in doc.tables:
            n_cols = len(table.columns)
            for col_idx in range(n_cols):
                for row in table.rows:
                    if col_idx < len(row.cells):
                        for para in row.cells[col_idx].paragraphs:
                            self._replace_in_paragraph(para, replacements)

        doc.save(str(output_path))

    @staticmethod
    def _iter_all_paragraphs(doc: "Document") -> Iterator[Any]:
        yield from doc.paragraphs
        for section in doc.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs

    @staticmethod
    def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
        full_text = "".join(r.text for r in paragraph.runs)
        if not full_text:
            return
        modified = full_text
        for original, placeholder in replacements.items():
            modified = modified.replace(original, placeholder)
        if modified == full_text:
            return
        # Склеиваем всё в первый run, остальные очищаем
        if paragraph.runs:
            paragraph.runs[0].text = modified
            for run in paragraph.runs[1:]:
                run.text = ""


class ExcelProcessor:
    """
    Extract/Rebuild для XLSX через openpyxl.

    Rebuild: замена во всех ячейках листа, сохранение формата.
    """

    def __init__(self) -> None:
        if not _OPENPYXL_OK:
            raise RuntimeError("openpyxl не установлен: pip install openpyxl")

    def extract_text(self, path: str | Path) -> str:
        parts: list[str] = []
        wb = load_workbook(str(path), read_only=True, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                parts.extend(str(cell) if cell is not None else "" for cell in row)
        wb.close()
        return "\n".join(parts)

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],
    ) -> None:
        wb = load_workbook(str(input_path))

        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        modified = cell.value
                        for original, placeholder in replacements.items():
                            modified = modified.replace(original, placeholder)
                        cell.value = modified

        wb.save(str(output_path))


class CSVProcessor:
    """
    Extract/Rebuild для CSV через pandas.

    Поддерживает различные кодировки и разделители.
    """

    def __init__(self) -> None:
        if not _PANDAS_OK:
            raise RuntimeError("pandas не установлен: pip install pandas")

    def detect_encoding(self, path: str | Path) -> str:
        """Определяет кодировку файла."""
        try:
            import chardet
            with open(path, 'rb') as f:
                result = chardet.detect(f.read())
            return result['encoding'] or 'utf-8'
        except ImportError:
            return 'utf-8'

    def detect_separator(self, path: str | Path, encoding: str) -> str:
        """Определяет разделитель (запятая, точка с запятой, таб)."""
        with open(path, 'r', encoding=encoding) as f:
            first_line = f.readline()
            separators = {',': first_line.count(','),
                         ';': first_line.count(';'),
                         '\t': first_line.count('\t')}
            return max(separators, key=separators.get)

    def extract_text(self, path: str | Path) -> str:
        encoding = self.detect_encoding(path)
        separator = self.detect_separator(path, encoding)
        df = pd.read_csv(path, encoding=encoding, sep=separator)

        # Собираем все значения в одну строку для детекции
        parts = []
        for col in df.columns:
            if df[col].dtype in ('object', 'str'):
                parts.extend(df[col].dropna().astype(str).tolist())
        return "\n".join(parts)

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],  # словарь замен
    ) -> None:
        encoding = self.detect_encoding(input_path)
        separator = self.detect_separator(input_path, encoding)
        df = pd.read_csv(input_path, encoding=encoding, sep=separator)

        # Анонимизируем все текстовые ячейки (поддерживает как 'object', так и 'str' типы)
        for col in df.columns:
            if df[col].dtype in ('object', 'str'):
                for original, placeholder in replacements.items():
                    df[col] = df[col].str.replace(original, placeholder, regex=False)

        df.to_csv(output_path, index=False, sep=separator, encoding=encoding)



class JSONProcessor:
    """
    Extract/Anonymize для JSON файлов.

    Рекурсивно обходит все строковые значения в JSON,
    анонимизирует их, сохраняет структуру.
    """

    def extract_text(self, path: str | Path) -> str:
        """Извлекает все строковые значения из JSON."""
        import json
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Извлекаем только строковые значения (без чисел)
        return self._extract_strings(data)

    def _extract_strings(self, obj: Any, separator: str = "\n") -> str:
        """Рекурсивно извлекает все строки из JSON."""
        if isinstance(obj, str):
            return obj
        elif isinstance(obj, list):
            return separator.join(self._extract_strings(item, separator) for item in obj if isinstance(item, (str, dict, list)))
        elif isinstance(obj, dict):
            return separator.join(self._extract_strings(v, separator) for v in obj.values() if isinstance(v, (str, dict, list)))
        return ""

    def _extract_pairs(self, obj: Any, separator: str = "\n") -> str:
        """Рекурсивно извлекает пары ключ-значение для детекции."""
        if isinstance(obj, str):
            return obj
        elif isinstance(obj, (int, float)):
            return str(obj)
        elif isinstance(obj, list):
            return separator.join(self._extract_pairs(item, separator) for item in obj)
        elif isinstance(obj, dict):
            parts = []
            for k, v in obj.items():
                if isinstance(v, (str, int, float)):
                    # Сохраняем пару ключ: значение для детекции
                    parts.append(f"{k}: {v}")
                elif isinstance(v, (dict, list)):
                    parts.append(self._extract_pairs(v, separator))
            return separator.join(parts)
        return ""

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],
    ) -> None:
        """Анонимизирует JSON, сохраняя структуру."""
        import json
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Собираем user_id/client_id отдельно
        user_id_map = self._extract_user_ids(data)

        # Рекурсивно анонимизируем с сохранением структуры
        anonymized_data = self._anonymize_obj(data, replacements, user_id_map)

        # Сохраняем с форматированием
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(anonymized_data, f, ensure_ascii=False, indent=2)

    def _extract_user_ids(self, obj: Any) -> dict[int, str]:
        """Извлекает user_id/client_id значений для замены."""
        user_ids = {}
        counter = 1  # Счетчик плейсхолдеров

        def extract_recursive(obj: Any) -> None:
            nonlocal counter
            if isinstance(obj, dict):
                for k, v in obj.items():
                    if k in ('user_id', 'client_id') and isinstance(v, int):
                        if v not in user_ids:
                            label = "ID_ПОЛЬЗОВАТЕЛЯ" if k == 'user_id' else "ID_КЛИЕНТА"
                            user_ids[v] = f"[{label}_{counter}]"
                            counter += 1
                    elif isinstance(v, (dict, list)):
                        extract_recursive(v)
            elif isinstance(obj, list):
                for item in obj:
                    extract_recursive(item)

        extract_recursive(obj)
        return user_ids

    def _anonymize_obj(self, obj: Any, replacements: dict[str, str], user_id_map: dict[int, str] | None = None) -> Any:
        """Рекурсивно анонимизирует все строки в объекте."""
        if user_id_map is None:
            user_id_map = {}

        if isinstance(obj, str):
            # Применяем ВСЕ замены к строке (не выходим после первой!)
            result = obj
            for original, placeholder in replacements.items():
                if original in result:
                    result = result.replace(original, placeholder)
            return result
        elif isinstance(obj, int):
            # Проверяем, есть ли это число в user_id_map
            return user_id_map.get(obj, obj)
        elif isinstance(obj, list):
            return [self._anonymize_obj(item, replacements, user_id_map) for item in obj]
        elif isinstance(obj, dict):
            return {k: self._anonymize_obj(v, replacements, user_id_map) for k, v in obj.items()}
        return obj


class ImageProcessor:
    """
    Extract/Anonymize для PNG/JPG изображений через Tesseract OCR.

    Извлекает текст с координатами, находит PII-сущности,
    замазывает области прямоугольниками.
    """

    def __init__(self) -> None:
        if not _PILLOW_OK:
            raise RuntimeError("Pillow не установлен: pip install Pillow")
        if not _TESSERACT_OK:
            raise RuntimeError("pytesseract не установлен: pip install pytesseract")

        # Проверяем наличие Tesseract
        try:
            pytesseract.get_tesseract_version()
        except pytesseract.TesseractNotFoundError:
            raise RuntimeError(
                "Tesseract не установлен.\n"
                "Ubuntu/Debian: sudo apt-get install tesseract-ocr tesseract-ocr-rus\n"
                "macOS: brew install tesseract tesseract-lang\n"
                "Windows: https://github.com/UB-Mannheim/tesseract/wiki"
            )

    def extract_text(self, path: str | Path, lang: str = "rus+eng") -> str:
        """Извлекает текст из изображения."""
        image = Image.open(str(path))
        return pytesseract.image_to_string(image, lang=lang)

    def _get_text_data(self, path: str | Path, lang: str = "rus+eng") -> list[dict]:
        """
        Извлекает текст с координатами.

        Возвращает список словарей:
        [
            {'text': 'Иванов', 'left': 10, 'top': 20, 'width': 50, 'height': 15},
            ...
        ]
        """
        image = Image.open(str(path))
        data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)

        results = []
        for i, text in enumerate(data['text']):
            if text.strip():  # Пропускаем пустые результаты
                results.append({
                    'text': text.strip(),
                    'left': data['left'][i],
                    'top': data['top'][i],
                    'width': data['width'][i],
                    'height': data['height'][i],
                })
        return results

    def anonymize_file(
        self,
        input_path:   str | Path,
        output_path:  str | Path,
        replacements: dict[str, str],
        fill_color:   tuple[int, int, int] = (0, 0, 0),
        lang: str = "rus+eng",
    ) -> None:
        """
        Замазывает PII-сущности на изображении черными прямоугольниками.

        replacements: словарь {оригинальный_текст: плейсхолдер}
        fill_color: цвет заливки (по умолчанию черный)
        lang: язык для Tesseract (rus+eng для русского и английского)
        """
        image = Image.open(str(input_path))
        draw = ImageDraw.Draw(image)

        # Получаем текст с координатами
        text_data = self._get_text_data(input_path, lang)

        # Находим и замазываем области с PII
        for item in text_data:
            text = item['text']
            for original, placeholder in replacements.items():
                if original in text:
                    # Вычисляем координаты прямоугольника
                    left = item['left']
                    top = item['top']
                    right = left + item['width']
                    bottom = top + item['height']

                    # Заливаем прямоугольник
                    draw.rectangle(
                        [left, top, right, bottom],
                        fill=fill_color
                    )
                    break  # Заменили - переходим к следующему

        # Сохраняем результат
        image.save(str(output_path))


# ══════════════════════════════════════════════════════════════
# § 9.  ГЛАВНЫЙ КЛАСС
# ══════════════════════════════════════════════════════════════

class Depersonalizer:
    """
    Оркестратор: Extract → Detect → Anonymize → Rebuild.

    Быстрый старт:
        dp = Depersonalizer(mode="placeholder")
        safe = dp.anonymize_text("Иванов И.И., тел. 8-900-123-45-67")
        dp.anonymize_file("input.pdf", "output_anon.pdf")
        dp.anonymize_file("input.docx", "output_anon.docx")
        dp.anonymize_file("input.xlsx", "output_anon.xlsx")
        dp.anonymize_file("input.csv", "output_anon.csv")
        dp.anonymize_file("data.json", "data_anon.json")
        dp.anonymize_file("screenshot.png", "screenshot_anon.png")
    """

    def __init__(
        self,
        mode:         str  = "placeholder",  # placeholder | hash | mask
        use_natasha:  bool = True,
        use_presidio: bool = True,
        use_stanza:   bool = False,
        language:     str  = "ru",
    ) -> None:
        self.tracker    = EntityTracker()
        self.detector   = EnsembleDetector(
            use_natasha=use_natasha,
            use_presidio=use_presidio,
            use_stanza=use_stanza,
            language=language,
        )
        self.anonymizer = TextAnonymizer(mode=mode, tracker=self.tracker)
        self._pdf_proc   = PDFProcessor()    if _PYMUPDF_OK   else None
        self._docx_proc  = DocxProcessor()   if _DOCX_OK      else None
        self._xlsx_proc  = ExcelProcessor()  if _OPENPYXL_OK  else None
        self._csv_proc   = CSVProcessor()    if _PANDAS_OK    else None
        self._json_proc  = JSONProcessor()   if True          else None  # JSON всегда доступен
        self._image_proc = ImageProcessor()  if _PILLOW_OK and _TESSERACT_OK else None

    # ── Public API ──────────────────────────────────────────

    def anonymize_text(self, text: str) -> str:
        """Detect + Anonymize для произвольной строки.

        Двухпроходная стратегия: после основной детекции ищем оставшиеся
        фрагменты найденных имён (фамилии без имени/отчества и т.п.).
        """
        matches = self.detector.detect(text)
        # Собираем слова из найденных PER/PERSON-сущностей для второго прохода
        name_words = {
            word
            for m in matches if m.entity_type in ("PER", "PERSON")
            for word in m.original_text.split()
            if len(word) >= 3  # пропускаем инициалы
        }
        # Первый проход анонимизации
        text = self.anonymizer.anonymize(text, matches)
        # Второй проход: ищем оставшиеся фрагменты имён
        if name_words:
            extra = [
                PIIMatch(
                    start=m.start(), end=m.end(),
                    entity_type="PERSON",
                    original_text=word,
                    confidence=0.80,
                    source="name_propagation",
                )
                for word in name_words
                for m in re.finditer(re.escape(word), text)
            ]
            if extra:
                text = self.anonymizer.anonymize(text, extra)
        return text

    def anonymize_file(
        self,
        input_path:  str | Path,
        output_path: str | Path | None = None,
    ) -> Path:
        """
        Extract → Detect → Anonymize → Rebuild для файла.
        Поддерживаемые форматы: .txt, .md, .log, .csv, .xlsx, .pdf, .docx, .json, .png, .jpg, .jpeg.
        Возвращает путь к анонимизированному файлу.
        """
        input_path  = Path(input_path)
        output_path = Path(output_path) if output_path else self._default_output(input_path)
        suffix = input_path.suffix.lower()

        if suffix in (".txt", ".md", ".log"):
            raw  = input_path.read_text(encoding="utf-8")
            anon = self.anonymize_text(raw)
            output_path.write_text(anon, encoding="utf-8")

        elif suffix == ".csv":
            if not self._csv_proc:
                raise RuntimeError("pandas не установлен: pip install pandas")
            raw = self._csv_proc.extract_text(input_path)
            self.anonymize_text(raw)  # заполняем tracker
            self._csv_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        elif suffix == ".xlsx":
            if not self._xlsx_proc:
                raise RuntimeError("openpyxl не установлен: pip install openpyxl")
            raw = self._xlsx_proc.extract_text(input_path)
            self.anonymize_text(raw)  # заполняем tracker
            self._xlsx_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        elif suffix == ".pdf":
            if not self._pdf_proc:
                raise RuntimeError("pymupdf не установлен: pip install pymupdf")
            raw = self._pdf_proc.extract_text(input_path)
            self.anonymize_text(raw)               # заполняем tracker
            self._pdf_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        elif suffix == ".docx":
            if not self._docx_proc:
                raise RuntimeError("python-docx не установлен: pip install python-docx")
            raw = self._docx_proc.extract_text(input_path)
            self.anonymize_text(raw)               # заполняем tracker
            self._docx_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        elif suffix == ".json":
            if not self._json_proc:
                raise RuntimeError("JSONProcessor недоступен")
            raw = self._json_proc.extract_text(input_path)
            self.anonymize_text(raw)  # заполняем tracker
            self._json_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        elif suffix in (".png", ".jpg", ".jpeg"):
            if not self._image_proc:
                raise RuntimeError("Pillow/pytesseract не установлены: pip install Pillow pytesseract")
            raw = self._image_proc.extract_text(input_path)
            self.anonymize_text(raw)               # заполняем tracker
            self._image_proc.anonymize_file(
                input_path, output_path, self.tracker.mapping()
            )

        else:
            raise ValueError(f"Неподдерживаемый формат: {suffix!r}")

        return output_path

    def get_report(self) -> dict[str, Any]:
        """Возвращает отчёт: общее кол-во сущностей + маппинг original→placeholder."""
        m = self.tracker.mapping()
        return {"total_entities": len(m), "mapping": m}

    @staticmethod
    def _default_output(p: Path) -> Path:
        return p.parent / f"{p.stem}_anonymized{p.suffix}"


# ══════════════════════════════════════════════════════════════
# § 10. CLI
# ══════════════════════════════════════════════════════════════

def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="Деперсонализация PII в файлах (PDF, DOCX, TXT, CSV, XLSX, JSON, PNG, JPG).",    
        )
    parser.add_argument(
        "file",
        type=Path,
        help="Путь к файлу для обработки (.txt, .pdf, .docx, .csv, .xlsx, .json, .png, .jpg)",
    )
    parser.add_argument(
        "-o", "--output",
        type=Path,
        default=None,
        help="Путь к выходному файлу (по умолчанию: <имя>_anonymized.<ext>)",
    )
    parser.add_argument(
        "-m", "--mode",
        choices=["placeholder", "hash", "mask"],
        default="placeholder",
        help="Режим анонимизации (по умолчанию: placeholder)",
    )
    parser.add_argument(
        "--no-natasha", action="store_true",
        help="Отключить Natasha NER",
    )
    parser.add_argument(
        "--no-presidio", action="store_true",
        help="Отключить Presidio NER",
    )
    parser.add_argument(
        "--use-stanza", action="store_true",
        help="Использовать Stanza вместо spaCy для Presidio",
    )
    parser.add_argument(
        "--lang",
        default="ru",
        help="Язык текста (по умолчанию: ru)",
    )

    args = parser.parse_args()

    if not args.file.exists():
        parser.error(f"Файл не найден: {args.file}")

    dp = Depersonalizer(
        mode=args.mode,
        use_natasha=not args.no_natasha,
        use_presidio=not args.no_presidio,
        use_stanza=args.use_stanza,
        language=args.lang,
    )

    result = dp.anonymize_file(args.file, args.output)
    report = dp.get_report()

    print(f"Готово: {result}")
    print(f"Найдено сущностей: {report['total_entities']}")
    if report["mapping"]:
        print("Замены:")
        for original, placeholder in report["mapping"].items():
            print(f"  {original!r} → {placeholder}")


if __name__ == "__main__":
    main()
